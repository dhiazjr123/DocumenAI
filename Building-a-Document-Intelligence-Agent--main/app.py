import streamlit as st
import os
import io
import tempfile
import requests
import time
import numpy as np
import re
from urllib.parse import urljoin
import pandas as pd
import difflib  # fuzzy matching nama kolom

# Mistral 0.4.2 (legacy client)
from mistralai.client import MistralClient

import google.generativeai as genai
from PIL import Image
from PyPDF2 import PdfReader

# --- Tambahan import untuk dashboard ---
import plotly.express as px
import plotly.graph_objects as go
from collections import Counter
import json
import string
import subprocess

# --- DOCX import with friendly error ---
try:
    # modul bernama "docx" disediakan oleh paket "python-docx"
    from docx import Document  # pip install python-docx
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except Exception as e:
    st.error(
        "Paket yang benar untuk .docx adalah **python-docx**. "
        "Sepertinya paket **docx** (yang salah) terpasang dan menyebabkan konflik.\n\n"
        "Perbaikan cepat:\n"
        "1) aktifkan venv\n"
        "2) `pip uninstall -y docx`\n"
        "3) `pip install --upgrade python-docx`\n\n"
        f"Detail error: {e}"
    )
    raise

# WordCloud opsional (fallback otomatis kalau belum terpasang)
try:
    from wordcloud import WordCloud
    HAS_WORDCLOUD = True
except Exception:
    HAS_WORDCLOUD = False


# --------------------- Page config ---------------------
st.set_page_config(page_title="Document Intelligence Agent", layout="wide")
st.title("Document Intelligence Agent")
st.caption("Upload documents or URL → extract, render (DOCX tables preserved), and ask questions.")

# --------------------- Global CSS (stabilkan tabel DOCX) ---------------------
st.markdown("""
<style>
.docx-table { width:100%; border-collapse:collapse; table-layout:fixed; margin:8px 0; }
.docx-table col { }
.docx-table th, .docx-table td {
  border:1px solid #ccc; padding:6px; vertical-align:top;
  white-space:pre-wrap; word-break:keep-all; overflow-wrap:anywhere;
}
.docx-table th { font-weight:700; }
</style>
""", unsafe_allow_html=True)

# --------------------- Sidebar: API Keys ----------------
with st.sidebar:
    st.header("API Configuration")
    mistral_api_key = st.text_input("Mistral AI API Key (legacy 0.4.2)", type="password")
    google_api_key = st.text_input("Google API Key (Gemini)", type="password")
    model_preference = st.selectbox(
        "Model preference",
        ["Auto (Pro→Flash)", "Flash only", "Pro only"],
        index=0,
        help="Fallback otomatis ke Flash saat Pro kena rate limit/kuota."
    )
    answer_language = st.selectbox(
        "Answer language",
        ["Bahasa Indonesia", "English"],
        index=0,
        help="Bahasa jawaban untuk fitur Q&A."
    )

    st.markdown("---")
    st.markdown("How To Get API Key Tutorials")
    st.markdown(
        """
- **Mistral AI API Key** — [YouTube Tutorial](https://youtu.be/NUCcUFwfhlA?si=iLrFFxVtcFUp657C)  
- **Google API Key (Gemini)** — [YouTube Tutorial](https://youtu.be/IHj7wF-8ry8?si=VKvhMM3pMeKwkXAv)
        """
    )

# Disimpan global agar helper bisa akses
MODEL_PREFERENCE = model_preference
ANSWER_LANGUAGE = answer_language

# MistralClient (legacy) — tidak untuk OCR di 0.4.2
mistral_client = None
if mistral_api_key:
    try:
        mistral_client = MistralClient(api_key=mistral_api_key)
        st.success("✅ Mistral API connected (legacy client 0.4.2)")
    except Exception as e:
        st.error(f"Failed to initialize Mistral client: {e}")

# Gemini untuk OCR & QnA
if google_api_key:
    try:
        genai.configure(api_key=google_api_key)
        st.success("✅ Google API connected")
    except Exception as e:
        st.error(f"Failed to initialize Google API: {e}")

# --------------------- Helpers (Gemini OCR) -------------------------
def _is_quota_error(err: Exception) -> bool:
    msg = str(err).lower()
    return "429" in msg or "quota" in msg or "rate limit" in msg or "exceeded" in msg

def _get_model_candidates() -> list:
    if MODEL_PREFERENCE == "Flash only":
        return ["gemini-1.5-flash"]
    if MODEL_PREFERENCE == "Pro only":
        return ["gemini-1.5-pro"]
    return ["gemini-1.5-pro", "gemini-1.5-flash"]

def _generate_with_fallback(parts_or_prompt):
    last_error = None
    candidates = _get_model_candidates()
    for model_name in candidates:
        try:
            model = genai.GenerativeModel(model_name=model_name)
            resp = model.generate_content(parts_or_prompt)
            text = getattr(resp, "text", "").strip()
            if text:
                if model_name != candidates[0]:
                    st.info(f"Using fallback model: {model_name}")
                return text
        except Exception as e:
            last_error = e
            if _is_quota_error(e):
                continue
            else:
                return f"Error generating response: {e}"
    if last_error and _is_quota_error(last_error):
        time.sleep(30)
        try:
            model = genai.GenerativeModel(model_name=candidates[-1])
            resp = model.generate_content(parts_or_prompt)
            return getattr(resp, "text", "").strip()
        except Exception as e2:
            return f"Error generating response: {e2}"
    return f"Error generating response: {last_error}"

def _truncate_context(text: str, max_chars: int = 20000) -> str:
    if not text:
        return text
    if len(text) <= max_chars:
        return text
    return text[:max_chars]

# --------------------- Text chunking & Embeddings (RAG) ---------------------
EMBEDDING_MODEL = "models/text-embedding-004"

def _chunk_text(text: str, chunk_size: int = 1800, overlap: int = 200) -> list:
    if not text:
        return []
    chunks = []
    start = 0
    end = max(chunk_size, 1)
    text_len = len(text)
    while start < text_len:
        chunk = text[start:end]
        chunks.append(chunk)
        if end >= text_len:
            break
        start = max(end - overlap, 0)
        end = min(start + chunk_size, text_len)
    return chunks

def _extract_embedding_values(resp) -> list:
    emb = getattr(resp, "embedding", None)
    if emb is not None:
        values = getattr(emb, "values", emb)
        if isinstance(values, (list, tuple)):
            return list(values)
        if isinstance(values, dict) and "values" in values:
            return list(values["values"])
    if isinstance(resp, dict):
        emb = resp.get("embedding")
        if isinstance(emb, dict) and "values" in emb:
            return list(emb["values"])
        if isinstance(emb, (list, tuple)):
            return list(emb)
    return []

def _embed_text(text: str) -> np.ndarray:
    try:
        resp = genai.embed_content(model=EMBEDDING_MODEL, content=text)
        values = _extract_embedding_values(resp)
        if not values:
            return np.array([])
        return np.array(values, dtype=np.float32)
    except Exception:
        return np.array([])

def _build_retrieval_index(full_text: str):
    text_hash = str(hash(full_text))
    if (st.session_state.get("cached_text_hash") == text_hash and 
        st.session_state.get("retrieval_embeddings") is not None):
        return
    chunks = _chunk_text(full_text)
    embeddings = []

    max_chunks = min(len(chunks), 10)
    chunks = chunks[:max_chunks]
    
    for ch in chunks:
        try:
            vec = _embed_text(ch)
            if vec.size == 0:
                embeddings = []
                break
            embeddings.append(vec)
        except Exception as e:
            if "429" in str(e) or "quota" in str(e).lower():
                st.warning("⚠️ Gemini embedding quota exceeded. Using full-context fallback.")
                embeddings = []
                break
            else:
                embeddings = []
                break
    
    if embeddings:
        matrix = np.vstack(embeddings)
        st.session_state.retrieval_chunks = chunks
        st.session_state.retrieval_embeddings = matrix
        st.session_state.retrieval_norms = np.linalg.norm(matrix, axis=1, keepdims=True) + 1e-10
        st.session_state.cached_text_hash = text_hash
    else:
        st.session_state.retrieval_chunks = None
        st.session_state.retrieval_embeddings = None
        st.session_state.retrieval_norms = None
        st.session_state.cached_text_hash = None

def _retrieve_top_k(query: str, k: int = 5) -> list:
    if not query:
        return []
    chunks = st.session_state.get("retrieval_chunks")
    emb = st.session_state.get("retrieval_embeddings")
    norms = st.session_state.get("retrieval_norms")
    if not chunks or emb is None or norms is None:
        return []
    q_vec = _embed_text(query)
    if q_vec.size == 0:
        return []
    q_vec = q_vec.reshape(1, -1)
    q_norm = np.linalg.norm(q_vec, axis=1, keepdims=True) + 1e-10
    sims = (emb @ q_vec.T) / (norms * q_norm)
    sims = sims.ravel()
    top_idx = np.argsort(-sims)[: max(1, k)]
    return [chunks[i] for i in top_idx]

def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        text = "\n".join(parts).strip()
        return text
    except Exception:
        return ""

# ===== DOCX → HTML (preserve tables, colspan/rowspan, basic styles) =====
from html import escape

def _run_to_html(run):
    text = escape(run.text or "")
    if not text:
        return ""
    if run.bold:
        text = f"<strong>{text}</strong>"
    if run.italic:
        text = f"<em>{text}</em>"
    if run.underline:
        text = f"<u>{text}</u>"
    return text

def _para_to_html(para):
    align = getattr(para.paragraph_format, "alignment", None)
    align_map = {0:"left", 1:"center", 2:"right", 3:"justify"}  # WD_ALIGN_PARAGRAPH
    style = ""
    if align is not None:
        style = f' style="text-align:{align_map.get(int(align), "left")};"'
    content = "".join(_run_to_html(r) for r in para.runs)
    return f"<p{style}>{content}</p>"

def _get_tblgrid_widths(table):
    """Ambil lebar kolom dari w:tblGrid/w:gridCol[@w:w] (dalam twips)."""
    try:
        grid = table._tbl.tblGrid
        if grid is None:
            max_cols = max(len(r.cells) for r in table.rows) if table.rows else 0
            return [1]*max_cols
        widths = []
        for gc in grid.iterchildren():
            if gc.tag.endswith('gridCol'):
                w = gc.get(qn('w:w'))
                widths.append(int(w) if w else 1)
        return widths if widths else [1]
    except Exception:
        max_cols = max(len(r.cells) for r in table.rows) if table.rows else 0
        return [1]*max_cols

def _cell_gridspan(cell) -> int:
    try:
        tcPr = cell._tc.tcPr
        if tcPr is not None and tcPr.gridSpan is not None:
            return int(tcPr.gridSpan.val)
    except Exception:
        pass
    return 1

def _cell_vmerge_state(cell) -> str | None:
    """
    return 'restart' | 'continue' | None
    """
    try:
        tcPr = cell._tc.tcPr
        if tcPr is None or tcPr.vMerge is None:
            return None
        v = tcPr.vMerge.val
        if v in (None, 'continue'):
            return 'continue'
        if v == 'restart':
            return 'restart'
        return 'continue'
    except Exception:
        return None

def _row_vrowspans(table):
    """
    Untuk setiap cell, kembalikan num rowspan efektif:
      - 0  → vMerge continue (jangan render)
      - >=1 → render di baris ini dengan rowspan tsb
    """
    n_rows = len(table.rows)
    result = []
    for r in range(n_rows):
        row = table.rows[r]
        arr = []
        for c, cell in enumerate(row.cells):
            stt = _cell_vmerge_state(cell)
            if stt == 'restart':
                span = 1
                rr = r+1
                while rr < n_rows:
                    try:
                        next_cell = table.rows[rr].cells[c]
                    except Exception:
                        break
                    st2 = _cell_vmerge_state(next_cell)
                    if st2 == 'continue':
                        span += 1; rr += 1
                    else:
                        break
                arr.append(span)
            elif stt == 'continue':
                arr.append(0)
            else:
                arr.append(1)
        result.append(arr)
    return result

def docx_table_to_html(table):
    """
    Render tabel DOCX ke HTML:
      - Menghormati w:tblGrid (lebar kolom) → <colgroup>
      - Menangani gridSpan (colspan) dan vMerge (rowspan) akurat
      - Skip sel yang tertutup merge (pakai matriks occupied pada grid kolom)
    """
    grid_widths = _get_tblgrid_widths(table)           # twips per grid column
    grid_cols = len(grid_widths)
    n_rows = len(table.rows)
    if n_rows == 0 or grid_cols == 0:
        return '<table class="docx-table"></table>'

    total_w = max(1, sum(grid_widths))
    col_perc = [max(1, int(round(w*100/total_w))) for w in grid_widths]
    diff = 100 - sum(col_perc)
    if diff != 0:
        col_perc[-1] = max(1, col_perc[-1] + diff)

    rowspans = _row_vrowspans(table)
    occupied = [[False]*grid_cols for _ in range(n_rows)]

    parts = []
    parts.append('<table class="docx-table">')
    parts.append('<colgroup>')
    for p in col_perc:
        parts.append(f'<col style="width:{p}%">')
    parts.append('</colgroup>')

    for r_idx, row in enumerate(table.rows):
        parts.append('<tr>')
        cpos = 0  # posisi grid kolom saat ini

        for c_idx, cell in enumerate(row.cells):
            rs = rowspans[r_idx][c_idx] if r_idx < len(rowspans) and c_idx < len(rowspans[r_idx]) else 1
            if rs == 0:
                continue  # vMerge-continue

            while cpos < grid_cols and occupied[r_idx][cpos]:
                cpos += 1
            if cpos >= grid_cols:
                continue

            cs = max(1, _cell_gridspan(cell))

            for rr in range(r_idx, min(n_rows, r_idx+rs)):
                for cc in range(cpos, min(grid_cols, cpos+cs)):
                    occupied[rr][cc] = True

            inner = "".join(_para_to_html(p) for p in cell.paragraphs) or "&nbsp;"
            tag = "th" if r_idx == 0 else "td"

            attrs = []
            if cs > 1:
                attrs.append(f'colspan="{cs}"')
            if rs > 1:
                attrs.append(f'rowspan="{rs}"')

            parts.append(f"<{tag} {' '.join(attrs)}>{inner}</{tag}>")
            cpos += cs

        parts.append('</tr>')

    parts.append('</table>')
    return "".join(parts)

def docx_to_html(doc):
    """Konversi DOCX ke HTML ringan dengan tabel dipertahankan."""
    out = []
    for block in doc.element.body.iterchildren():
        tag = block.tag
        if tag == qn('w:tbl'):
            tbl = None
            for t in doc.tables:
                if t._tbl is block:
                    tbl = t; break
            if tbl is not None:
                out.append(docx_table_to_html(tbl))
        elif tag == qn('w:p'):
            for para in doc.paragraphs:
                if para._p is block:
                    out.append(_para_to_html(para))
                    break
    return "\n".join(out).strip()

# --------------------- DOC/DOCX extractors -------------------------
def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """Return HTML yang mempertahankan struktur tabel (colspan/rowspan, align, bold/italic)."""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(docx_bytes)
            tmp_path = tmp.name
        doc = Document(tmp_path)

        html = docx_to_html(doc)

        if not html or len(html) < 20:
            parts = []
            parts.extend(_para_to_html(p) for p in doc.paragraphs if p.text)
            for tbl in doc.tables:
                parts.append(docx_table_to_html(tbl))
            html = "\n".join(parts)

        return html
    except Exception:
        return ""
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass

def extract_text_from_doc_bytes(doc_bytes: bytes) -> str:
    try:
        with tempfile.TemporaryDirectory() as td:
            src = os.path.join(td, "in.doc")
            with open(src, "wb") as f:
                f.write(doc_bytes)

            # unoconv -> txt
            try:
                subprocess.run(["unoconv", "-f", "txt", src], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                txt_path = src.replace(".doc", ".txt")
                if os.path.exists(txt_path):
                    with open(txt_path, "r", encoding="utf-8", errors="ignore") as t:
                        return t.read()
            except Exception:
                pass

            # soffice -> txt
            try:
                subprocess.run(
                    ["soffice", "--headless", "--convert-to", "txt:Text", "--outdir", td, src],
                    check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE
                )
                base = os.path.splitext(os.path.basename(src))[0]
                txt_path = os.path.join(td, base + ".txt")
                if os.path.exists(txt_path):
                    with open(txt_path, "r", encoding="utf-8", errors="ignore") as t:
                        return t.read()
            except Exception:
                pass

            # soffice -> pdf -> extractor PDF / OCR Gemini
            try:
                subprocess.run(
                    ["soffice", "--headless", "--convert-to", "pdf", "--outdir", td, src],
                    check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE
                )
                pdf_path = os.path.join(td, "in.pdf")
                if os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as pf:
                        pdf_bytes = pf.read()
                    txt = extract_text_from_pdf_bytes(pdf_bytes)
                    if txt and len(txt) > 30:
                        return txt
                    return gemini_ocr_pdf(pdf_bytes, filename="converted_from_doc.pdf")
            except Exception:
                pass

            # strings (last resort)
            try:
                out = subprocess.run(["strings", src], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                txt = out.stdout.decode("utf-8", errors="ignore")
                return txt
            except Exception:
                return ""
    except Exception:
        return ""

# --------------------- OCR helpers -------------------------
def gemini_ocr_image(image_bytes: bytes) -> str:
    img = Image.open(io.BytesIO(image_bytes))
    prompt_doc = (
        "Convert this document image into clean Markdown. "
        "Preserve headings, lists, and tables (use Markdown tables). "
        "Maintain natural reading order."
    )
    text = _generate_with_fallback([prompt_doc, img])
    if not text or len(text.strip()) < 30:
        prompt_cap_id = (
            "Jelaskan gambar ini secara ringkas, jelas, dan akurat. "
            "Sebutkan objek utama, konteks, warna, teks (jika ada), dan hal penting lainnya."
        )
        prompt_cap_en = (
            "Describe this image concisely and accurately. "
            "Mention main objects, context, colors, any visible text, and other important details."
        )
        prompt_cap = prompt_cap_id if ANSWER_LANGUAGE == "Bahasa Indonesia" else prompt_cap_en
        text = _generate_with_fallback([prompt_cap, img])
    return text

def gemini_ocr_pdf(pdf_bytes: bytes, filename: str = "upload.pdf") -> str:
    if not google_api_key:
        return "Please provide a valid Google API Key for OCR/processing."
    suffix = os.path.splitext(filename)[1] or ".pdf"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(pdf_bytes)
        tmp_path = tmp.name
    try:
        file_obj = genai.upload_file(
            path=tmp_path,
            mime_type="application/pdf",
            display_name=filename
        )
        try:
            for _ in range(30):
                f = genai.get_file(file_obj.name)
                state = getattr(getattr(f, "state", None), "name", getattr(f, "state", ""))
                if str(state).upper() == "ACTIVE":
                    break
                time.sleep(1)
            else:
                return "Error: File processing timed out. Please try again."
        except Exception:
            time.sleep(2)
        prompt = (
            "Extract the full content of this PDF as clean Markdown. "
            "Preserve headings and tables. If pages are scanned, perform OCR first."
        )
        return _generate_with_fallback([file_obj, prompt])
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass

# --------------------- PDF/Image/DOC/DOCX pipeline -------------------------
def process_document_with_gemini(kind: str, name: str, data: bytes) -> str:
    if kind == "pdf":
        raw_text = extract_text_from_pdf_bytes(data)
        looks_like_table = bool(re.search(r'\|\s*[^|]+\s*\|', raw_text))
        if len(raw_text) >= 200 and looks_like_table:
            return raw_text
        return gemini_ocr_pdf(data, filename=name)

    if kind == "image":
        return gemini_ocr_image(data)

    if kind == "docx":
        html = extract_text_from_docx_bytes(data)
        if html and len(html) >= 50:
            return html
        return _generate_with_fallback([
            "Extract the full content of this Office document as HTML with tables preserved.",
            data
        ])

    if kind == "doc":
        text = extract_text_from_doc_bytes(data)
        if text and len(text) >= 50:
            return text
        return "No content extracted from .doc (older Word). Please ensure LibreOffice/unoconv is installed for better results."

    return "Unsupported document kind."

def answer_from_image(image_bytes: bytes, question: str) -> str:
    try:
        img = Image.open(io.BytesIO(image_bytes))
        if ANSWER_LANGUAGE == "Bahasa Indonesia":
            prompt = (
                "Anda adalah asisten analisis visual. Jawab pertanyaan pengguna hanya berdasarkan gambar ini.\n"
                "Jika informasi tidak terlihat pada gambar, katakan tidak ada.\n"
                f"Pertanyaan: {question}"
            )
        else:
            prompt = (
                "You are a visual analysis assistant. Answer the user's question based only on this image.\n"
                "If the information is not visible in the image, say so.\n"
                f"Question: {question}"
            )
        return _generate_with_fallback([prompt, img]) or "No response text."
    except Exception as e:
        return f"Error generating visual answer: {e}"

def generate_response(context: str, query: str) -> str:
    if not context or len(context) < 10:
        if "image_bytes" in st.session_state and isinstance(st.session_state.image_bytes, dict) and st.session_state.image_bytes:
            first_img = next(iter(st.session_state.image_bytes.values()))
            return answer_from_image(first_img, query)
    try:
        retrieved_chunks = _retrieve_top_k(query, k=5)
        if retrieved_chunks:
            context_block = "\n\n---\n\n".join(retrieved_chunks)
        else:
            context_block = context
        
        doc_context = ""
        if st.session_state.get("documents") and len(st.session_state.documents) > 1:
            doc_names = [doc['name'] for doc in st.session_state.documents]
            doc_context = f"\n\nAvailable documents: {', '.join(doc_names)}"
        
        if ANSWER_LANGUAGE == "Bahasa Indonesia":
            prompt = f"""
Anda adalah asisten analisis dokumen. Gunakan konteks berikut untuk menjawab:

{context_block}{doc_context}

Pertanyaan pengguna:
{query}

Jawab dalam Bahasa Indonesia secara ringkas, jelas, dan akurat. Jika jawabannya tidak terdapat pada konteks, katakan: "Tidak ditemukan pada dokumen". Jika ada beberapa dokumen, sebutkan dari dokumen mana informasi berasal.
"""
        else:
            prompt = f"""
You are a document analysis assistant. Use only the context below to answer:

{context_block}{doc_context}

User question:
{query}

Respond in English concisely. If the answer is not in the context, say so. If there are multiple documents, mention which document the information comes from.
"""
        return _generate_with_fallback(prompt) or "No response text."
    except Exception as e:
        return f"Error generating response: {e}"

# --------------------- API Fallback & Caching ---------------------
def generate_with_mistral_fallback(prompt: str) -> str:
    if not mistral_client:
        return "Error: No Mistral API key configured for fallback."
    try:
        response = mistral_client.chat(
            model="mistral-large-latest",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error with Mistral fallback: {e}"

def generate_response_with_fallback(context: str, query: str) -> str:
    try:
        return generate_response(context, query)
    except Exception as e:
        if "429" in str(e) or "quota" in str(e).lower():
            st.warning("⚠️ Gemini API quota exceeded. Falling back to Mistral API...")
            if ANSWER_LANGUAGE == "Bahasa Indonesia":
                mistral_prompt = f"""
Anda adalah asisten analisis dokumen. Gunakan konteks berikut untuk menjawab:

{context}

Pertanyaan pengguna:
{query}

Jawab dalam Bahasa Indonesia secara ringkas, jelas, dan akurat. Jika jawabannya tidak terdapat pada konteks, katakan: "Tidak ditemukan pada dokumen".
"""
            else:
                mistral_prompt = f"""
You are a document analysis assistant. Use only the context below to answer:

{context}

User question:
{query}

Respond in English concisely. If the answer is not in the context, say so.
"""
            return generate_with_mistral_fallback(mistral_prompt)
        else:
            return f"Error generating response: {e}"

# --------------------- Document Management Helpers ---------------------
def clear_all_document_state():
    st.session_state.documents = []
    st.session_state.ocr_content = None
    st.session_state.retrieval_chunks = None
    st.session_state.retrieval_embeddings = None
    st.session_state.retrieval_norms = None
    st.session_state.image_bytes = {}
    st.session_state.chat_history = []

def rebuild_document_content():
    if st.session_state.documents:
        all_content = [f"--- DOCUMENT: {d['name']} ---\n{d['content']}" for d in st.session_state.documents]
        st.session_state.ocr_content = "\n\n".join(all_content)
        _build_retrieval_index(st.session_state.ocr_content)
    else:
        st.session_state.ocr_content = None
        st.session_state.retrieval_chunks = None
        st.session_state.retrieval_embeddings = None
        st.session_state.retrieval_norms = None


# ======================= DASHBOARD HELPERS =======================
_ID_STOPWORDS = {
    "yang","dan","di","ke","dari","untuk","pada","dengan","ini","itu","ada","tidak","atau","karena","sebagai",
    "dalam","atas","oleh","sebuah","para","akan","juga","sudah","belum","saat","kami","kita","mereka","anda",
    "ia","dia","tersebut","rp","usd","pt","tbk","persero","co","ltd","inc"
}
_EN_STOPWORDS = {
    "the","and","of","to","in","for","on","at","by","with","from","as","is","are","was","were","be","been","a","an",
    "this","that","these","those","it","its","we","you","they","he","she","them","our","your","their",
    "or","not","but","if","then","so","than","such","per","vs"
}
STOPWORDS = _ID_STOPWORDS | _EN_STOPWORDS

def _tokenize(text: str) -> list:
    text = text.lower()
    text = text.translate(str.maketrans({c: " " for c in string.punctuation}))
    toks = [t for t in text.split() if t and t not in STOPWORDS and not t.isdigit() and len(t) > 2]
    return toks

def _compute_ocr_quality(text: str) -> float:
    if not text: return 0.0
    n = len(text)
    good = sum(ch.isalnum() or ch.isspace() or ch in ".,:;-%()[]|/+\n" for ch in text)
    bad = text.count(" ")
    short_lines = sum(1 for ln in text.splitlines() if 0 < len(ln.strip()) < 3)
    score = (good / n) * 100.0
    score -= min(25, bad * 0.5)
    score -= min(15, short_lines * 0.2)
    return max(0.0, min(100.0, score))

def _structure_stats(md_text: str) -> dict:
    if not md_text:
        return {"text": 0, "tables": 0, "images": 0}
    lines = md_text.splitlines()
    table_lines = sum(1 for ln in lines if "<table" in md_text or (ln.strip().startswith("|") and ln.count("|") >= 2))
    image_tags = md_text.count("![") + md_text.lower().count("<img")
    text_chars = len(md_text)
    return {"text": text_chars, "tables": table_lines, "images": image_tags}

def _extract_sections(md_text: str):
    sections = []
    current_title = "Intro"
    current_buf = []
    for ln in md_text.splitlines():
        if ln.startswith("--- DOCUMENT:"):
            if current_buf:
                sections.append((current_title, "\n".join(current_buf).strip()))
                current_buf = []
            current_title = ln.replace("--- DOCUMENT:", "").strip()
        elif re.match(r"^\s{0,3}#{1,3}\s+\S", ln):
            if current_buf:
                sections.append((current_title, "\n".join(current_buf).strip()))
                current_buf = []
            current_title = re.sub(r"^\s{0,3}#{1,3}\s+", "", ln).strip()
        else:
            current_buf.append(ln)
    if current_buf:
        sections.append((current_title, "\n".join(current_buf).strip()))
    return sections[:12] if sections else [("All", md_text)]

def _missing_matrix(sections):
    cats = ["N/A/NA", "Empty Lines", "Dashes(-/—)", "Question(?)"]
    labels = [title[:40] + ("…" if len(title) > 40 else "") for title,_ in sections]
    matrix = []
    for _, txt in sections:
        lines = txt.splitlines()
        na = sum(bool(re.search(r"\b(n/?a|tidak tersedia|kosong)\b", ln, re.I)) for ln in lines)
        empty = sum(1 for ln in lines if not ln.strip())
        dashes = sum(ln.count("-") + ln.count("—") for ln in lines)
        ques = txt.count("?")
        matrix.append([na, empty, dashes, ques])
    return labels, cats, np.array(matrix).T

def _is_financial_report(text: str) -> bool:
    keys = ["revenue","pendapatan","penjualan","income","profit","laba","rugi",
            "neraca","balance sheet","arus kas","cash flow","laba kotor","gross","operating","net"]
    t = text.lower()
    return any(k in t for k in keys)

_num_pat = re.compile(r"([\-]?\d{1,3}(?:[\.,]\d{3})*(?:[\.,]\d{1,2})?)")

def _parse_number(s: str):
    s = s.strip()
    if not s: return None
    if "." in s and "," in s:
        if s.find(".") < s.find(","):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
    else:
        if "," in s and "." not in s:
            s = s.replace(",", ".")
    try:
        return float(s)
    except Exception:
        digits = re.sub(r"[^\d\-\.]", "", s)
        try:
            return float(digits)
        except Exception:
            return None

def _extract_financials(text: str):
    rev = {}
    prof = {}
    gross = {}
    op = {}
    net = {}
    assets = equity = curr_assets = curr_liab = debt = None

    lines = text.splitlines()
    for ln in lines:
        years = re.findall(r"\b(20\d{2}|19\d{2})\b", ln)
        if not years:
            continue
        y = int(years[0])

        if re.search(r"\b(revenue|pendapatan|penjualan)\b", ln, re.I):
            m = _num_pat.search(ln)
            if m:
                val = _parse_number(m.group(1))
                if val is not None:
                    rev[y] = val
        if re.search(r"\b(net income|laba bersih|laba/rugi bersih|profit|laba)\b", ln, re.I):
            m = _num_pat.search(ln)
            if m:
                val = _parse_number(m.group(1))
                if val is not None:
                    prof[y] = val
        if re.search(r"\b(gross profit|laba kotor)\b", ln, re.I):
            m = _num_pat.search(ln)
            if m:
                val = _parse_number(m.group(1))
                if val is not None:
                    gross[y] = val
        if re.search(r"\b(operating profit|laba usaha|laba operasi)\b", ln, re.I):
            m = _num_pat.search(ln)
            if m:
                val = _parse_number(m.group(1))
                if val is not None:
                    op[y] = val
        if re.search(r"\b(net profit|net income|laba bersih)\b", ln, re.I):
            m = _num_pat.search(ln)
            if m:
                val = _parse_number(m.group(1))
                if val is not None:
                    net[y] = val

        if assets is None and re.search(r"\b(total assets|jumlah aset)\b", ln, re.I):
            m = _num_pat.search(ln); assets = _parse_number(m.group(1)) if m else None
        if equity is None and re.search(r"\b(total equity|ekuitas)\b", ln, re.I):
            m = _num_pat.search(ln); equity = _parse_number(m.group(1)) if m else None
        if curr_assets is None and re.search(r"\b(current assets|aset lancar)\b", ln, re.I):
            m = _num_pat.search(ln); curr_assets = _parse_number(m.group(1)) if m else None
        if curr_liab is None and re.search(r"\b(current liab|liabilitas lancar|utang lancar)\b", ln, re.I):
            m = _num_pat.search(ln); curr_liab = _parse_number(m.group(1)) if m else None
        if debt is None and re.search(r"\b(total debt|utang|pinjaman)\b", ln, re.I):
            m = _num_pat.search(ln); debt = _parse_number(m.group(1)) if m else None

    return {
        "revenue_by_year": dict(sorted(rev.items())),
        "profit_by_year": dict(sorted(prof.items())),
        "gross_by_year": dict(sorted(gross.items())),
        "operating_by_year": dict(sorted(op.items())),
        "net_by_year": dict(sorted(net.items())),
        "assets": assets, "equity": equity,
        "current_assets": curr_assets, "current_liabilities": curr_liab, "debt": debt
    }

def _yoy_growth(series: dict) -> dict:
    ys = sorted(series.keys())
    growth = {}
    for i in range(1, len(ys)):
        y0, y1 = ys[i-1], ys[i]
        if series[y0] and series[y0] != 0:
            growth[y1] = (series[y1] - series[y0]) / abs(series[y0]) * 100.0
    return growth

def _readability_from_avg_sentence_len(text: str):
    if not text: return 0.0, 0.0
    sents = re.split(r"[.!?\n]+", text)
    sents = [s.strip() for s in sents if s.strip()]
    if not sents: return 0.0, 0.0
    words = sum(len(s.split()) for s in sents)
    avg = words / len(sents)
    score = 100.0 - ((avg - 8) / (40 - 8)) * 100.0
    score = max(0.0, min(100.0, score))
    return avg, score

def _ner_counts_with_gemini(text: str):
    if not google_api_key:
        return None
    try:
        prompt = """
Extract counts of named entities by coarse type from the text.
Only return a compact JSON like {"ORG": 12, "PERSON": 5, "LOC": 7}. Types limited to ORG, PERSON, LOC.
If none detected, use 0.
Text:
""" + _truncate_context(text, 20000)
        raw = _generate_with_fallback(prompt)
        data = json.loads(re.findall(r"\{.*\}", raw, re.S)[0])
        return {"ORG": int(data.get("ORG", 0)), "PERSON": int(data.get("PERSON", 0)), "LOC": int(data.get("LOC", 0))}
    except Exception:
        return None

def _ner_counts_naive(text: str):
    ORG = len(re.findall(r"\b(PT|Tbk|Persero|Inc\.?|Ltd\.?|LLC|Corp\.?)\b", text))
    PERSON = len(re.findall(r"\b[A-Z][a-z]+ [A-Z][a-z]+(?: [A-Z][a-z]+)?\b", text))
    LOC = len(re.findall(r"\b(Jakarta|Bandung|Surabaya|Medan|Indonesia|Singapore|Malaysia|USA|Europe|Asia)\b", text))
    return {"ORG": ORG, "PERSON": PERSON, "LOC": LOC}


# ======================= TABLE PARSER & AVG INTENT =======================
_TABLE_ROW_RE = re.compile(r'^\s*\|(.+)\|\s*$')
_TABLE_SEP_RE = re.compile(r'^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$')

def _parse_markdown_tables(md_text: str) -> list:
    """
    Mengembalikan list of (DataFrame, meta) dari tabel Markdown di teks.
    Meta: {"doc_section": str, "table_index": int}
    Mencoba infer header saat tidak ada header eksplisit.
    """
    tables = []
    lines = md_text.splitlines()
    buf = []

    def _flush(buf_lines, t_idx):
        if not buf_lines:
            return
        rows = []
        for ln in buf_lines:
            m = _TABLE_ROW_RE.match(ln)
            if not m:
                continue
            cells = [c.strip() for c in m.group(1).split('|')]
            rows.append(cells)

        if not rows:
            return

        header = None
        if len(rows) >= 2 and _TABLE_SEP_RE.match(buf_lines[1]):  # header | sep | data...
            header = [h if h else f"col_{i+1}" for i, h in enumerate(rows[0])]
            data = rows[2:] if len(rows) > 2 else []
        else:
            first = rows[0]
            nonnum = sum(1 for x in first if not re.fullmatch(r'[-+]?\d+([.,]\d+)?', (x or '')))
            if nonnum >= max(1, len(first) // 2):
                header = [h if h else f"col_{i+1}" for i, h in enumerate(first)]
                data = rows[1:]
            else:
                ncol = len(first)
                header = [f"col_{i+1}" for i in range(ncol)]
                data = rows

        maxc = len(header)
        norm_rows = []
        for r in data:
            if len(r) < maxc:
                r = r + [''] * (maxc - len(r))
            elif len(r) > maxc:
                r = r[:maxc]
            norm_rows.append(r)

        try:
            df = pd.DataFrame(norm_rows, columns=header)
            tables.append((df, {"doc_section": "All", "table_index": t_idx, "source": "md"}))
        except Exception:
            pass

    t_idx = 0
    for ln in lines:
        if _TABLE_ROW_RE.match(ln):
            buf.append(ln)
        else:
            if buf:
                _flush(buf, t_idx)
                t_idx += 1
                buf = []
    if buf:
        _flush(buf, t_idx)

    return tables

# === NEW: parser tabel HTML (hasil DOCX) ===
def _parse_html_tables(html_text: str) -> list:
    """
    Parse <table> HTML (hasil DOCX) menjadi list of (DataFrame, meta).
    Prioritas: pandas.read_html; fallback BeautifulSoup jika perlu.
    """
    tables = []
    if not html_text or "<table" not in html_text.lower():
        return tables

    # 1) Coba langsung dengan pandas
    try:
        dfs = pd.read_html(html_text)  # type: ignore
        for i, df in enumerate(dfs):
            cols = [str(c) if str(c).strip() else f"col_{j+1}" for j, c in enumerate(df.columns)]
            df.columns = cols
            tables.append((df, {"doc_section": "All", "table_index": i, "source": "html"}))
        if tables:
            return tables
    except Exception:
        pass

    # 2) Fallback ringan pakai BeautifulSoup
    try:
        from bs4 import BeautifulSoup  # pip install beautifulsoup4
        soup = BeautifulSoup(html_text, "html.parser")
        for i, tbl in enumerate(soup.find_all("table")):
            rows = []
            for tr in tbl.find_all("tr"):
                cells = [c.get_text(" ", strip=True) for c in tr.find_all(["th", "td"])]
                if cells:
                    rows.append(cells)
            if not rows:
                continue
            maxc = max(len(r) for r in rows)
            norm = [r + [""]*(maxc-len(r)) for r in rows]
            first_tr = tbl.find_all("tr")[0] if tbl.find_all("tr") else None
            use_header = first_tr and all(x.name == "th" for x in first_tr.find_all(["th","td"]))
            if use_header:
                header, data = norm[0], norm[1:]
            else:
                header = [f"col_{j+1}" for j in range(maxc)]
                data = norm
            df = pd.DataFrame(data, columns=header)
            tables.append((df, {"doc_section": "All", "table_index": i, "source": "html"}))
    except Exception:
        pass

    return tables

def _is_avg_intent(q: str) -> bool:
    ql = q.lower()
    keys = ["rata rata", "rata-rata", "average", "avg", "mean"]
    return any(k in ql for k in keys)

def _is_visualization_intent(q: str) -> bool:
    """Detect if user wants a visualization/chart."""
    ql = q.lower()
    viz_keywords = [
        "grafik", "chart", "visualisasi", "visualization", "plot", "diagram", 
        "bar chart", "line chart", "pie chart", "histogram", "scatter plot",
        "tampilkan grafik", "buat grafik", "show chart", "create chart",
        "gambar grafik", "draw chart", "plot data", "visualize data"
    ]
    return any(keyword in ql for keyword in viz_keywords)

def _extract_target_column_terms(q: str) -> list:
    quoted = re.findall(r'[""](.+?)[""]', q)
    terms = [t.strip() for t in quoted if t.strip()]
    m = re.search(r'\b(kolom|column|field)\s+([A-Za-z0-9_\-\s]+)', q, re.I)
    if m:
        chunk = m.group(2).strip()
        chunk = re.split(r'[,.;:?]|rata|average|mean', chunk, 1)[0].strip()
        if chunk:
            terms.append(chunk)
    toks = _tokenize(q)
    toks = [t for t in toks if t not in {"rata", "rata-rata", "average", "mean", "avg", "nilai", "value"}]
    if toks:
        terms.append(max(toks, key=len))
    seen = set(); res = []
    for t in terms:
        tt = t.lower()
        if tt not in seen:
            seen.add(tt); res.append(t)
    return res[:3]

def _normalize_col(s: str) -> str:
    s = s.strip().lower()
    s = re.sub(r'\s+', ' ', s)
    s = s.replace('_', ' ')
    return s

def _guess_best_column(df: pd.DataFrame, query: str) -> str | None:
    if df.empty or df.shape[1] == 0:
        return None
    colmap = {c: _normalize_col(c) for c in df.columns}
    terms = _extract_target_column_terms(query)
    synonyms = {
        "harga": ["price", "amount", "nilai", "nominal", "jumlah"],
        "pendapatan": ["revenue", "sales", "omzet"],
        "laba": ["profit", "net profit", "laba bersih", "income"],
        "tanggal": ["date", "periode", "period", "bulan", "month", "year", "tahun"],
        "qty": ["quantity", "jumlah", "kuantitas", "volume"],
        # penting untuk kasus nilai mahasiswa
        "nilai": ["score", "skor", "uts", "uas", "nilai uts", "nilai uas", "nilai uts/uas", "nilai ujian"],
    }
    candidates = []
    for t in terms:
        nt = _normalize_col(t)
        for orig, norm in colmap.items():
            if nt == norm or nt in norm or norm in nt:
                candidates.append(orig)
        close = difflib.get_close_matches(nt, list(colmap.values()), n=1, cutoff=0.75)
        if close:
            for orig, norm in colmap.items():
                if norm == close[0]:
                    candidates.append(orig)
        for base, syns in synonyms.items():
            if nt == base or nt in syns:
                close2 = difflib.get_close_matches(base, list(colmap.values()), n=1, cutoff=0.6)
                if close2:
                    for orig, norm in colmap.items():
                        if norm == close2[0]:
                            candidates.append(orig)
    for c in candidates:
        s = pd.to_numeric(
            df[c].astype(str)
                .str.replace(r'[^\d\-\.,]', '', regex=True)
                .str.replace('.', '', regex=False)
                .str.replace(',', '.', regex=False),
            errors='coerce'
        )
        if s.notna().sum() >= max(2, int(len(s)*0.5)):
            return c
    best_col = None; best_score = -1
    for c in df.columns:
        s = pd.to_numeric(
            df[c].astype(str)
                .str.replace(r'[^\d\-\.,]', '', regex=True)
                .str.replace('.', '', regex=False)
                .str.replace(',', '.', regex=False),
            errors='coerce'
        )
        score = s.notna().sum()
        if score > best_score:
            best_score = score; best_col = c
    return best_col

# --------- Fallback parser nilai dari teks polos ---------
def _to_float_id_en(num: str) -> float | None:
    if not num:
        return None
    s = num.strip()
    if "." in s and "," in s:
        if s.find(".") < s.find(","):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
    else:
        if "," in s and "." not in s:
            s = s.replace(",", ".")
    try:
        return float(s)
    except Exception:
        try:
            digits = re.sub(r"[^\d\.\-]", "", s)
            return float(digits) if digits else None
        except Exception:
            return None

def _extract_scores_from_text(text: str) -> list[float]:
    scores: list[float] = []
    t = " ".join(text.split())
    p_phone = re.compile(r'(\d{1,3}(?:[.,]\d{1,2})?)\s+(0\d{9,13})')
    for m in p_phone.finditer(t):
        val = _to_float_id_en(m.group(1))
        if val is not None and 0 <= val <= 100:
            scores.append(val)
    p_rank = re.compile(
        r'\b\d{1,3}\s+(?:\d{8,12})\s+[A-Za-zÀ-ÿ\.\'\- ]{2,}?'
        r'\s+(\d{1,3}(?:[.,]\d{1,2})?)\b'
    )
    for m in p_rank.finditer(t):
        val = _to_float_id_en(m.group(1))
        if val is not None and 0 <= val <= 100:
            scores.append(val)
    if not scores:
        near = []
        for m in re.finditer(r'(nilai[^0-9]{0,30})(\d{1,3}(?:[.,]\d{1,2})?)', t, flags=re.I):
            val = _to_float_id_en(m.group(2))
            if val is not None and 0 <= val <= 100:
                near.append(val)
        if len(near) >= 3:
            scores.extend(near)
    if scores:
        seen = set(); uniq = []
        for v in scores:
            key = round(v, 2)
            if key not in seen:
                seen.add(key); uniq.append(v)
        scores = uniq
    return scores

# ========== BARU: pilih dokumen berdasar pertanyaan ==========
def _normalize_filename(name: str) -> str:
    base = os.path.splitext(name)[0]
    base = re.sub(r'[^a-zA-Z0-9]+', ' ', base).strip().lower()
    return base

def select_docs_for_query(query: str, documents: list) -> list:
    q = query.lower()
    matches = []
    for d in documents:
        norm = _normalize_filename(d["name"])
        if norm and norm in q:
            matches.append(d)
        else:
            tokens = [t for t in norm.split() if len(t) > 2]
            hit = sum(1 for t in tokens if t in q)
            if hit >= max(1, len(tokens)//2):
                matches.append(d)
    if len(matches) == 1:
        return matches
    if len(matches) >= 2:
        return matches
    return documents

# ========== MODIFIED: hitung rata-rata per dokumen (HTML + Markdown) ==========
def compute_average_per_doc(query: str, documents: list) -> list[dict]:
    if not _is_avg_intent(query):
        return []

    results = []
    for doc in documents:
        content = doc.get("content") or ""
        tables = []

        # DOCX → HTML tables
        if "<table" in content.lower():
            tables += _parse_html_tables(content)

        # PDF/Markdown tables
        tables += _parse_markdown_tables(content)

        best = None

        for df, meta in tables:
            if df.empty:
                continue
            col = _guess_best_column(df, query)
            if not col:
                continue
            s = pd.to_numeric(
                df[col].astype(str)
                    .str.replace(r'[^\d\-\.,]', '', regex=True)
                    .str.replace('.', '', regex=False)   # ribuan → buang titik
                    .str.replace(',', '.', regex=False),  # desimal id → '.'
                errors='coerce'
            )
            vals = s.dropna()
            # singkirkan angka tidak masuk akal (mis. no. telepon)
            vals = vals[(vals >= 0) & (vals <= 100)]
            if len(vals) == 0:
                continue

            mean_val = float(vals.mean())
            cand = {"doc": doc["name"], "column": col, "value": mean_val, "n": int(len(vals))}
            if (best is None) or (cand["n"] > best["n"]):
                best = cand

        # Fallback teks polos tetap ada
        if (best is None) or best["n"] < 3:
            scores = _extract_scores_from_text(content)
            if len(scores) >= 3:
                mean_val = float(np.mean(scores))
                best = {"doc": doc["name"], "column": "Nilai (fallback)", "value": mean_val, "n": int(len(scores))}

        if best:
            results.append(best)

    return results

def generate_visualization(query: str, documents: list) -> tuple[bool, str, any]:
    """
    Generate visualization based on user query and document data.
    Returns: (success, message, chart_object)
    """
    if not _is_visualization_intent(query):
        return False, "", None
    
    try:
        # Debug info
        st.info(f"🔍 Processing visualization request: {query}")
        st.info(f"📚 Found {len(documents)} documents")
        
        # Extract all tables from documents
        all_tables = []
        for doc in documents:
            content = doc.get("content") or ""
            
            # DOCX → HTML tables
            if "<table" in content.lower():
                html_tables = _parse_html_tables(content)
                all_tables.extend(html_tables)
                st.info(f"📊 Found {len(html_tables)} HTML tables in {doc.get('name', 'Unknown')}")
            
            # PDF/Markdown tables
            md_tables = _parse_markdown_tables(content)
            all_tables.extend(md_tables)
            st.info(f"📊 Found {len(md_tables)} Markdown tables in {doc.get('name', 'Unknown')}")
        
        if not all_tables:
            st.warning("⚠️ No tables found in documents to visualize.")
            return False, "No tables found in documents to visualize.", None
        
        st.success(f"✅ Total tables found: {len(all_tables)}")
        
        # Find the best table for visualization
        best_table = None
        best_score = 0
        
        for i, (df, meta) in enumerate(all_tables):
            if df.empty or df.shape[0] < 2 or df.shape[1] < 2:
                continue
            
            # Score based on data quality
            numeric_cols = 0
            text_cols = 0
            for col in df.columns:
                try:
                    # Try to convert to numeric
                    pd.to_numeric(df[col].astype(str).str.replace(r'[^\d\-\.,]', '', regex=True), errors='coerce')
                    numeric_cols += 1
                except:
                    text_cols += 1
            
            score = numeric_cols * 2 + text_cols
            st.info(f"📈 Table {i+1}: {df.shape[0]} rows × {df.shape[1]} cols, {numeric_cols} numeric, {text_cols} text, score: {score}")
            
            if score > best_score:
                best_score = score
                best_table = (df, meta)
        
        if not best_table:
            st.warning("⚠️ No suitable table data found for visualization.")
            return False, "No suitable table data found for visualization.", None
        
        df, meta = best_table
        st.success(f"✅ Selected best table: {df.shape[0]} rows × {df.shape[1]} cols")
        
        # Determine chart type based on query
        ql = query.lower()
        chart_type = "auto"
        
        if any(x in ql for x in ["bar", "bar chart", "grafik batang"]):
            chart_type = "bar"
        elif any(x in ql for x in ["line", "line chart", "grafik garis", "trend"]):
            chart_type = "line"
        elif any(x in ql for x in ["pie", "pie chart", "grafik lingkaran", "donut"]):
            chart_type = "pie"
        elif any(x in ql for x in ["scatter", "scatter plot", "scatterplot", "scatter plot"]):
            chart_type = "scatter"
        elif any(x in ql for x in ["histogram", "distribusi", "distribution"]):
            chart_type = "histogram"
        
        st.info(f"🎯 Chart type detected: {chart_type}")
        
        # Generate appropriate chart
        chart = _create_chart(df, chart_type, query, documents)
        
        if chart:
            st.success(f"✅ Successfully generated {chart_type} chart!")
            return True, f"Generated {chart_type} chart from {meta.get('doc_section', 'document')} data.", chart
        else:
            st.error(f"❌ Failed to generate {chart_type} chart.")
            return False, "Failed to generate chart. Please try with different data or chart type.", None
            
    except Exception as e:
        st.error(f"❌ Error generating visualization: {str(e)}")
        import traceback
        st.error(f"Traceback: {traceback.format_exc()}")
        return False, f"Error generating visualization: {str(e)}", None

def _create_document_comparison_data(query: str, documents: list) -> pd.DataFrame:
    """Create aggregated data for document comparison in pie charts."""
    try:
        if not documents or len(documents) < 2:
            return None
            
        comparison_data = []
        
        for doc in documents:
            try:
                doc_name = doc.get("name", "Unknown")
                content = doc.get("content", "")
                
                # Count rows in tables or estimate from content
                table_count = 0
                
                # Method 1: Count table rows from HTML
                if "<table" in content.lower():
                    # Count table rows more accurately
                    table_rows = content.lower().count("<tr>")
                    table_headers = content.lower().count("<th>")
                    table_count = max(1, table_rows - table_headers)  # Subtract header rows
                
                # Method 2: Count from markdown tables
                elif "|" in content:
                    lines = content.split('\n')
                    table_lines = [line for line in lines if '|' in line and line.strip()]
                    # Remove separator lines (like |---|---|)
                    separator_lines = [line for line in table_lines if re.match(r'^\s*\|[\s\-:]+\|\s*$', line)]
                    table_count = max(1, len(table_lines) - len(separator_lines))
                
                # Method 3: Estimate from content length and structure
                else:
                    lines = content.split('\n')
                    # Look for patterns that suggest data rows
                    data_lines = [line for line in lines if len(line.strip()) > 10 and 
                                (re.search(r'\d+', line) or re.search(r'[A-Za-z]+\s+[A-Za-z]+', line))]
                    table_count = max(1, len(data_lines))
                
                # Ensure we have a reasonable count
                if table_count <= 0:
                    table_count = 1
                
                comparison_data.append({
                    'document': doc_name,
                    'count': table_count
                })
                
            except Exception as e:
                # If individual document processing fails, continue with others
                st.warning(f"Warning: Could not process document {doc.get('name', 'Unknown')}: {str(e)}")
                continue
        
        if len(comparison_data) >= 2:  # Need at least 2 documents to compare
            df_result = pd.DataFrame(comparison_data)
            # Sort by count descending for better visualization
            df_result = df_result.sort_values('count', ascending=False)
            return df_result
        
        return None
        
    except Exception as e:
        st.error(f"Error creating document comparison data: {str(e)}")
        return None

def _create_chart(df: pd.DataFrame, chart_type: str, query: str, documents: list) -> any:
    """Create a specific type of chart from DataFrame."""
    try:
        if df.empty or df.shape[0] < 2:
            return None
        
        # Clean and prepare data
        df_clean = df.copy()
        
        # Find numeric columns
        numeric_cols = []
        for col in df.columns:
            try:
                pd.to_numeric(df_clean[col].astype(str).str.replace(r'[^\d\-\.,]', '', regex=True), errors='coerce')
                numeric_cols.append(col)
            except:
                continue
        
        if not numeric_cols:
            return None
        
        # Smart column selection based on query context
        value_col = _select_best_value_column(df_clean, query, numeric_cols)
        if not value_col:
            value_col = numeric_cols[0]  # fallback
        
        # Clean numeric data
        df_clean[value_col] = pd.to_numeric(
            df_clean[value_col].astype(str)
                .str.replace(r'[^\d\-\.,]', '', regex=True)
                .str.replace('.', '', regex=False)
                .str.replace(',', '.', regex=False),
            errors='coerce'
        )
        
        # Remove rows with NaN values
        df_clean = df_clean.dropna(subset=[value_col])
        
        if df_clean.empty or len(df_clean) < 2:
            return None
        
        # Limit data points for better visualization
        if len(df_clean) > 20:
            df_clean = df_clean.head(20)
        
        # Create chart based on type
        if chart_type == "bar":
            # Find a good label column
            label_col = _select_best_label_column(df_clean, query, value_col)
            
            if label_col:
                fig = px.bar(
                    df_clean, 
                    x=label_col, 
                    y=value_col,
                    title=f"Grafik Batang: {value_col} berdasarkan {label_col}",
                    labels={value_col: "Nilai", label_col: "Kategori"}
                )
            else:
                # Use index as labels
                fig = px.bar(
                    df_clean, 
                    x=df_clean.index, 
                    y=value_col,
                    title=f"Grafik Batang: {value_col}",
                    labels={value_col: "Nilai", "index": "Index"}
                )
        
        elif chart_type == "line":
            # Try to find a sequential column (like dates, years, etc.)
            seq_col = _select_best_sequence_column(df_clean, query, value_col)
            
            if seq_col:
                df_clean[seq_col] = pd.to_numeric(df_clean[seq_col], errors='coerce')
                df_clean = df_clean.sort_values(seq_col)
                fig = px.line(
                    df_clean, 
                    x=seq_col, 
                    y=value_col,
                    title=f"Grafik Garis: {value_col} berdasarkan {seq_col}",
                    labels={value_col: "Nilai", seq_col: "Urutan"}
                )
            else:
                fig = px.line(
                    df_clean, 
                    x=df_clean.index, 
                    y=value_col,
                    title=f"Grafik Garis: {value_col}",
                    labels={value_col: "Nilai", "index": "Index"}
                )
        
        elif chart_type == "pie":
            # Find a good label column
            label_col = _select_best_label_column(df_clean, query, value_col)
            
            if label_col:
                fig = px.pie(
                    df_clean, 
                    values=value_col, 
                    names=label_col,
                    title=f"Grafik Lingkaran: {value_col} berdasarkan {label_col}"
                )
            else:
                # Try to create pie chart for document comparison
                if "membandingkan" in query.lower() or "vs" in query.lower() or "antara" in query.lower():
                    try:
                        # Create aggregated data for document comparison
                        doc_comparison_data = _create_document_comparison_data(query, documents)
                        if doc_comparison_data is not None and not doc_comparison_data.empty:
                            fig = px.pie(
                                doc_comparison_data,
                                values='count',
                                names='document',
                                title="Grafik Lingkaran: Perbandingan Jumlah Kandidat Antar Dokumen"
                            )
                        else:
                            return None
                    except Exception as e:
                        st.error(f"Error creating document comparison pie chart: {str(e)}")
                        return None
                else:
                    # Try to create pie chart from existing data by creating categories
                    try:
                        categorized_data = _create_categorized_pie_data(df_clean, value_col, query)
                        if categorized_data is not None and not categorized_data.empty:
                            fig = px.pie(
                                categorized_data,
                                values='count',
                                names='category',
                                title=f"Grafik Lingkaran: Distribusi {value_col} berdasarkan Kategori"
                            )
                        else:
                            return None
                    except Exception as e:
                        st.error(f"Error creating categorized pie chart: {str(e)}")
                        return None
        
        elif chart_type == "scatter":
            # Need two numeric columns
            if len(numeric_cols) >= 2:
                x_col = _select_best_x_column(df_clean, query, numeric_cols, value_col)
                df_clean[x_col] = pd.to_numeric(
                    df_clean[x_col].astype(str)
                        .str.replace(r'[^\d\-\.,]', '', regex=True)
                        .str.replace('.', '', regex=False)
                        .str.replace(',', '.', regex=False),
                    errors='coerce'
                )
                df_clean = df_clean.dropna(subset=[x_col])
                
                if not df_clean.empty and len(df_clean) >= 2:
                    fig = px.scatter(
                        df_clean, 
                        x=x_col, 
                        y=value_col,
                        title=f"Scatter Plot: {value_col} vs {x_col}",
                        labels={value_col: "Nilai Y", x_col: "Nilai X"}
                    )
                else:
                    return None
            else:
                return None  # Scatter needs two numeric columns
        
        elif chart_type == "histogram":
            fig = px.histogram(
                df_clean, 
                x=value_col,
                title=f"Histogram: Distribusi {value_col}",
                labels={value_col: "Nilai"},
                nbins=min(20, len(df_clean))
            )
        
        else:  # auto - try to pick the best
            if len(numeric_cols) >= 2:
                # Try scatter plot
                chart = _create_chart(df, "scatter", query, documents)
                if chart:
                    return chart
            
            # Try bar chart
            chart = _create_chart(df, "bar", query, documents)
            if chart:
                return chart
            
            # Fallback to histogram
            return _create_chart(df, "histogram", query, documents)
        
        # Update layout for better appearance
        fig.update_layout(
            height=500,
            showlegend=True,
            margin=dict(l=50, r=50, t=80, b=50),
            font=dict(size=12),
            hovermode='closest'
        )
        
        # Add some interactivity
        fig.update_traces(
            hovertemplate='<b>%{x}</b><br>Nilai: %{y}<extra></extra>'
        )
        
        # Add unique identifier to prevent duplicate element ID errors
        fig.update_layout(
            title=dict(text=fig.layout.title.text + f" (ID: {hash(str(fig))})")
        )
        
        return fig
        
    except Exception as e:
        st.error(f"Error creating chart: {str(e)}")
        return None

def _select_best_value_column(df: pd.DataFrame, query: str, numeric_cols: list) -> str:
    """Select the best column for values based on query context."""
    query_lower = query.lower()
    
    # Priority keywords for value columns
    priority_keywords = {
        'uts': ['uts', 'ujian tengah semester', 'midterm'],
        'uas': ['uas', 'ujian akhir semester', 'final', 'ujian akhir'],
        'nilai': ['nilai', 'score', 'skor', 'grade', 'hasil'],
        'pendapatan': ['pendapatan', 'revenue', 'income', 'penjualan'],
        'harga': ['harga', 'price', 'cost', 'biaya'],
        'kuantitas': ['kuantitas', 'quantity', 'jumlah', 'volume']
    }
    
    # Check for specific column names first
    for col in df.columns:
        col_lower = col.lower()
        for keyword_group, keywords in priority_keywords.items():
            if any(keyword in col_lower for keyword in keywords):
                if col in numeric_cols:
                    return col
    
    # Check for keywords in query
    for keyword_group, keywords in priority_keywords.items():
        if any(keyword in query_lower for keyword in keywords):
            # Find best matching column
            for col in numeric_cols:
                col_lower = col.lower()
                if any(keyword in col_lower for keyword in keywords):
                    return col
    
    # If no specific match, return first numeric column
    return numeric_cols[0] if numeric_cols else None

def _select_best_label_column(df: pd.DataFrame, query: str, value_col: str) -> str:
    """Select the best column for labels based on query context."""
    query_lower = query.lower()
    
    # Priority keywords for label columns
    label_keywords = {
        'kandidat': ['kandidat', 'candidate', 'nama', 'name', 'mahasiswa', 'student'],
        'bulan': ['bulan', 'month', 'periode', 'period'],
        'kategori': ['kategori', 'category', 'jenis', 'type', 'kelompok']
    }
    
    # Check for specific column names first
    for col in df.columns:
        if col == value_col:
            continue
        col_lower = col.lower()
        
        # Skip numeric columns for labels
        try:
            pd.to_numeric(df[col], errors='coerce')
            continue
        except:
            pass
        
        # Check if column has reasonable number of unique values
        unique_count = df[col].nunique()
        if unique_count > 20:  # Too many unique values
            continue
            
        for keyword_group, keywords in label_keywords.items():
            if any(keyword in col_lower for keyword in keywords):
                return col
    
    # Check for keywords in query
    for keyword_group, keywords in label_keywords.items():
        if any(keyword in query_lower for keyword in keywords):
            for col in df.columns:
                if col == value_col:
                    continue
                col_lower = col.lower()
                try:
                    pd.to_numeric(df[col], errors='coerce')
                    continue
                except:
                    pass
                if df[col].nunique() <= 20:
                    return col
    
    # Find best non-numeric column with reasonable unique values
    for col in df.columns:
        if col == value_col:
            continue
        try:
            pd.to_numeric(df[col], errors='coerce')
            continue
        except:
            pass
        if df[col].nunique() <= 20:
            return col
    
    return None

def _select_best_sequence_column(df: pd.DataFrame, query: str, value_col: str) -> str:
    """Select the best column for sequence/x-axis in line charts."""
    query_lower = query.lower()
    
    # Priority keywords for sequence columns
    sequence_keywords = {
        'urutan': ['urutan', 'order', 'ranking', 'peringkat', 'rank'],
        'waktu': ['waktu', 'time', 'tanggal', 'date', 'bulan', 'month', 'tahun', 'year'],
        'index': ['index', 'nomor', 'number', 'id']
    }
    
    # Check for specific column names first
    for col in df.columns:
        if col == value_col:
            continue
        col_lower = col.lower()
        
        for keyword_group, keywords in sequence_keywords.items():
            if any(keyword in col_lower for keyword in keywords):
                try:
                    pd.to_numeric(df[col], errors='coerce')
                    return col
                except:
                    continue
    
    # Check for keywords in query
    for keyword_group, keywords in sequence_keywords.items():
        if any(keyword in query_lower for keyword in keywords):
            for col in df.columns:
                if col == value_col:
                    continue
                try:
                    pd.to_numeric(df[col], errors='coerce')
                    return col
                except:
                    continue
    
    # Return first numeric column that's not the value column
    for col in df.columns:
        if col != value_col:
            try:
                pd.to_numeric(df[col], errors='coerce')
                return col
            except:
                continue
    
    return None

def _select_best_x_column(df: pd.DataFrame, query: str, numeric_cols: list, value_col: str) -> str:
    """Select the best column for x-axis in scatter plots."""
    # Return first numeric column that's not the value column
    for col in numeric_cols:
        if col != value_col:
            return col
    return numeric_cols[0] if len(numeric_cols) > 1 else None

def _create_categorized_pie_data(df: pd.DataFrame, value_col: str, query: str) -> pd.DataFrame:
    """Create categorized data for pie charts by grouping numeric values into ranges."""
    try:
        if df.empty or value_col not in df.columns:
            return None
        
        # Get numeric values
        values = pd.to_numeric(df[value_col], errors='coerce').dropna()
        if len(values) < 2:
            return None
        
        # Create categories based on value ranges
        categories = []
        counts = []
        
        # For scores (0-100), create grade categories
        if values.max() <= 100 and values.min() >= 0:
            # Grade categories
            grade_ranges = [
                (90, 100, "A (90-100)"),
                (80, 89, "B (80-89)"),
                (70, 79, "C (70-79)"),
                (60, 69, "D (60-69)"),
                (0, 59, "E (<60)")
            ]
            
            for min_val, max_val, label in grade_ranges:
                count = len(values[(values >= min_val) & (values <= max_val)])
                if count > 0:
                    categories.append(label)
                    counts.append(count)
        
        # For other numeric data, create quartile categories
        else:
            try:
                q25, q50, q75 = values.quantile([0.25, 0.5, 0.75])
                
                quartile_ranges = [
                    (values.min(), q25, f"Q1 ({values.min():.1f}-{q25:.1f})"),
                    (q25, q50, f"Q2 ({q25:.1f}-{q50:.1f})"),
                    (q50, q75, f"Q3 ({q50:.1f}-{q75:.1f})"),
                    (q75, values.max(), f"Q4 ({q75:.1f}-{values.max():.1f})")
                ]
                
                for min_val, max_val, label in quartile_ranges:
                    count = len(values[(values >= min_val) & (values <= max_val)])
                    if count > 0:
                        categories.append(label)
                        counts.append(count)
            except Exception:
                # Fallback: create simple ranges
                min_val, max_val = values.min(), values.max()
                range_size = (max_val - min_val) / 4
                
                for i in range(4):
                    start = min_val + (i * range_size)
                    end = min_val + ((i + 1) * range_size) if i < 3 else max_val
                    count = len(values[(values >= start) & (values <= end)])
                    if count > 0:
                        categories.append(f"Range {i+1} ({start:.1f}-{end:.1f})")
                        counts.append(count)
        
        if len(categories) > 1:  # Need at least 2 categories
            return pd.DataFrame({
                'category': categories,
                'count': counts
            })
        
        return None
        
    except Exception as e:
        st.error(f"Error creating categorized pie data: {str(e)}")
        return None

# --------------------- UI Layout -----------------------
col1, col2 = st.columns([1, 1])

with col1:
    st.header("Document Upload")
    uploaded_files = st.file_uploader(
        "Upload multiple documents (DOC, DOCX, PDF, PNG, JPG, JPEG)",
        type=["pdf", "png", "jpg", "jpeg", "doc", "docx"],
        accept_multiple_files=True,
    )
    url_input = st.text_input("Or enter a URL (web page or document):")
    st.session_state["url_input"] = url_input

    process_button = st.button("Process Documents")

    if "documents" not in st.session_state:
        st.session_state.documents = []
    if "ocr_content" not in st.session_state:
        st.session_state.ocr_content = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "show_chart_examples" not in st.session_state:
        st.session_state.show_chart_examples = False

    # --------------------- URL processing helper -----------------------
    def process_url_to_content(url: str) -> tuple:
        try:
            r = requests.get(
                url,
                timeout=30,
                headers={
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
                },
                allow_redirects=True,
            )
            r.raise_for_status()
            content_type = r.headers.get("Content-Type", "").lower()
            clean_url = url.split("?")[0]
            ext = os.path.splitext(clean_url)[1].lower()
            data = r.content

            # signature
            is_pdf_sig = data[:4] == b"%PDF"
            is_png_sig = data[:8] == b"\x89PNG\r\n\x1a\n"
            is_jpg_sig = data[:3] == b"\xff\xd8\xff"

            # Detect Office CT
            is_docx_ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in content_type
            is_doc_ct = "application/msword" in content_type

            chosen_kind = None
            if is_pdf_sig or "pdf" in content_type or ext == ".pdf":
                chosen_kind = "pdf"
            elif is_png_sig or is_jpg_sig or any(img_ct in content_type for img_ct in ["image/png", "image/jpeg", "image/jpg"]) or ext in [".png", ".jpg", ".jpeg"]:
                chosen_kind = "image"
            elif ext == ".docx" or is_docx_ct:
                chosen_kind = "docx"
            elif ext == ".doc" or is_doc_ct:
                chosen_kind = "doc"

            if not chosen_kind and content_type.startswith("text/html"):
                html_text = r.text
                links = re.findall(
                    r'href=["\']([^"\']+\.(?:pdf|png|jpe?g|docx?|DOCX?))(?:[#\?][^"\']*)?["\']',
                    html_text, flags=re.IGNORECASE
                )
                if links:
                    target_url = urljoin(url, links[0])
                    rr = requests.get(
                        target_url,
                        timeout=30,
                        headers={
                            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
                        },
                        allow_redirects=True,
                    )
                    rr.raise_for_status()
                    target_ct = rr.headers.get("Content-Type", "").lower()
                    tdata = rr.content
                    t_clean = target_url.split("?")[0]
                    t_ext = os.path.splitext(t_clean)[1].lower()

                    if tdata[:4] == b"%PDF" or "pdf" in target_ct or t_ext == ".pdf":
                        chosen_kind = "pdf"; data = tdata; clean_url = t_clean
                    elif tdata[:8] == b"\x89PNG\r\n\x1a\n" or tdata[:3] == b"\xff\xd8\xff" or any(ic in target_ct for ic in ["image/png","image/jpeg","image/jpg"]) or t_ext in [".png",".jpg",".jpeg"]:
                        chosen_kind = "image"; data = tdata; clean_url = t_clean
                        st.session_state.image_bytes = data
                    elif t_ext == ".docx" or "vnd.openxmlformats-officedocument.wordprocessingml.document" in target_ct:
                        chosen_kind = "docx"; data = tdata; clean_url = t_clean
                    elif t_ext == ".doc" or "application/msword" in target_ct:
                        chosen_kind = "doc"; data = tdata; clean_url = t_clean

                if not chosen_kind:
                    stripped = re.sub(r"<script[\s\S]*?</script>", " ", html_text, flags=re.IGNORECASE)
                    stripped = re.sub(r"<style[\s\S]*?</style>", " ", stripped, flags=re.IGNORECASE)
                    text_only = re.sub(r"<[^>]+>", " ", stripped)
                    text_only = re.sub(r"\s+", " ", text_only).strip()
                    st.session_state.ocr_content = text_only
                    if st.session_state.ocr_content:
                        _build_retrieval_index(st.session_state.ocr_content)
                        return True, "Webpage processed as text."
                    return False, "No content extracted from webpage."

            if not chosen_kind:
                return False, f"Unsupported content type: {content_type or ext}"

            st.session_state.ocr_content = process_document_with_gemini(
                chosen_kind, os.path.basename(clean_url) or "download", data
            )
            if chosen_kind == "image":
                if "image_bytes" not in st.session_state or not isinstance(st.session_state.image_bytes, dict):
                    st.session_state.image_bytes = {}
                st.session_state.image_bytes[os.path.basename(clean_url) or "image"] = data

            if st.session_state.ocr_content:
                _build_retrieval_index(st.session_state.ocr_content)
                return True, "Document processed successfully!"
            return False, "No content extracted."
        except Exception as e:
            return False, f"Error processing document: {e}"

    if process_button:
        if not google_api_key:
            st.error("Please provide a valid Google API Key for OCR/processing.")
        if uploaded_files:
            with st.spinner("Processing documents..."):
                all_content = []
                for uploaded_file in uploaded_files:
                    try:
                        ext = os.path.splitext(uploaded_file.name)[1].lower()
                        if ext == ".pdf":
                            kind = "pdf"
                        elif ext in [".png", ".jpg", ".jpeg"]:
                            kind = "image"
                        elif ext == ".docx":
                            kind = "docx"
                        elif ext == ".doc":
                            kind = "doc"
                        else:
                            kind = "unknown"

                        if kind == "image":
                            if "image_bytes" not in st.session_state:
                                st.session_state.image_bytes = {}
                            st.session_state.image_bytes[uploaded_file.name] = uploaded_file.getvalue()
                        
                        if kind == "unknown":
                            st.warning(f"Unsupported file type for {uploaded_file.name}. Skipped.")
                            continue

                        content = process_document_with_gemini(
                            kind, uploaded_file.name, uploaded_file.getvalue()
                        )
                        
                        if content:
                            doc_info = {
                                "name": uploaded_file.name,
                                "type": kind,
                                "content": content,
                                "size": len(uploaded_file.getvalue())
                            }
                            st.session_state.documents.append(doc_info)
                            all_content.append(f"--- DOCUMENT: {uploaded_file.name} ---\n{content}")
                            st.success(f"Document {uploaded_file.name} processed successfully!")
                        else:
                            st.warning(f"No content extracted from {uploaded_file.name}.")
                    except Exception as e:
                        st.error(f"Error processing {uploaded_file.name}: {e}")
                
                if all_content:
                    st.session_state.ocr_content = "\n\n".join(all_content)
                    _build_retrieval_index(st.session_state.ocr_content)
                    st.success(f"All {len(uploaded_files)} documents processed and combined!")
        if url_input:
            with st.spinner("Downloading & processing from URL..."):
                success, msg = process_url_to_content(url_input)
                if success:
                    if st.session_state.get("ocr_content"):
                        url_doc_info = {
                            "name": f"URL: {url_input[:50]}...",
                            "type": "url",
                            "content": st.session_state.ocr_content,
                            "size": len(st.session_state.ocr_content)
                        }
                        st.session_state.documents.append(url_doc_info)
                    st.success(msg)
                else:
                    st.error(msg)
        if not uploaded_files and not url_input:
            st.warning("Please upload a document or provide a URL.")

with col2:
    st.header("Document Q&A")

    if st.session_state.documents:
        st.markdown(f"**{len(st.session_state.documents)} document(s) loaded.**")
        
        total_chars = sum(doc['size'] for doc in st.session_state.documents)
        total_mb = total_chars / (1024 * 1024)
        
        c1, c2, c3 = st.columns([2, 2, 1])
        with c1:
            st.metric("Total Documents", len(st.session_state.documents))
        with c2:
            st.metric("Total Content", f"{total_chars:,} chars")
        with c3:
            st.metric("Memory Usage", f"{total_mb:.1f} MB")
        
        if len(st.session_state.documents) > 10:
            st.warning("⚠️ Many documents loaded. Consider removing some to improve performance.")
        elif total_chars > 500000:
            st.warning("⚠️ Very large document collection. Consider removing some documents to avoid processing limits.")
        elif total_chars > 200000 and len(st.session_state.documents) > 3:
            st.warning("⚠️ Large document collection. Consider removing some documents to avoid processing limits.")
        
        with st.expander("📋 Document List"):
            for i, doc in enumerate(st.session_state.documents):
                cc1, cc2 = st.columns([4, 1])
                with cc1:
                    st.markdown(f"**{i+1}. {doc['name']}** ({doc['type']}) - {doc['size']} chars")
                with cc2:
                    if st.button(f"🗑️", key=f"del_{i}", help=f"Delete {doc['name']}"):
                        deleted_doc = st.session_state.documents.pop(i)
                        if deleted_doc['type'] == 'image' and 'image_bytes' in st.session_state:
                            if isinstance(st.session_state.image_bytes, dict):
                                st.session_state.image_bytes.pop(deleted_doc['name'], None)
                            else:
                                st.session_state.image_bytes = {}
                        rebuild_document_content()
                        st.rerun()
            
            if st.session_state.documents:
                if st.button("🔄 Reset All Documents", type="secondary"):
                    clear_all_document_state()
                    st.rerun()
        
        # ---------- Quick actions ----------
        if st.session_state.documents:
            st.markdown("**Quick Actions:**")
            q1, q2, q3 = st.columns(3)
            with q1:
                if st.button("🗑️ Clear Chat", help="Clear chat history but keep documents"):
                    st.session_state.chat_history = []
                    st.rerun()
            with q2:
                if st.button("📊 Document Stats", help="Show dashboard with charts"):
                    st.session_state.show_stats = not st.session_state.get('show_stats', False)
                    st.rerun()
            with q3:
                if st.button("📈 Chart Examples", help="Show examples of chart requests"):
                    st.session_state.show_chart_examples = not st.session_state.get('show_chart_examples', False)
                    st.rerun()
        
        # ---------- Chart Examples ----------
        if st.session_state.get('show_chart_examples', False):
            st.markdown("### 📈 **Contoh Permintaan Grafik**")
            st.markdown("""
            **Grafik Batang (Bar Chart):**
            - "Buat grafik batang dari data penjualan"
            - "Tampilkan grafik batang nilai UTS/UAS kandidat"
            - "Buat grafik batang perbandingan pendapatan antar bulan"
            - "Buat grafik batang untuk membandingkan nilai antar dokumen"
            
            **Grafik Garis (Line Chart):**
            - "Tampilkan grafik garis pendapatan dari waktu ke waktu"
            - "Buat grafik garis trend nilai siswa"
            - "Buat grafik garis untuk melihat perkembangan nilai"
            - "Generate grafik garis trend data"
            
            **Grafik Lingkaran (Pie Chart):**
            - "Buat grafik lingkaran dari pangsa pasar"
            - "Tampilkan grafik lingkaran distribusi nilai"
            - "Buat grafik lingkaran persentase kandidat"
            - "Show me grafik lingkaran data"
            
            **Scatter Plot:**
            - "Generate scatter plot harga vs kuantitas"
            - "Buat scatter plot nilai UTS vs UAS"
            - "Buat scatter plot korelasi dua variabel"
            - "Show me scatter plot data"
            
            **Histogram:**
            - "Buat histogram dari nilai ujian"
            - "Tampilkan histogram distribusi data"
            - "Buat histogram untuk melihat sebaran nilai"
            - "Show me histogram data"
            
            **Auto-detection:** Sistem akan otomatis memilih tipe grafik terbaik jika Anda hanya meminta "grafik" atau "visualisasi".
            """)

        # ---------- DASHBOARD ----------
        if st.session_state.get('show_stats', False):
            st.markdown("## 📊 Document Analytics Dashboard")

            # Data agregat per dokumen
            stats_data = []
            combined_texts = []
            struct_total = {"text":0, "tables":0, "images":0}
            for doc in st.session_state.documents:
                txt = doc['content'] or ""
                combined_texts.append(txt)
                ocr_q = _compute_ocr_quality(txt)
                words = len(_tokenize(txt))
                lines = len(txt.splitlines())
                stt = {
                    "Document": doc['name'],
                    "Type": doc['type'],
                    "Size (chars)": doc['size'],
                    "Words": words,
                    "Lines": lines,
                    "OCR Quality": round(ocr_q,1)
                }
                stats_data.append(stt)
                s = _structure_stats(txt)
                struct_total["text"] += s["text"]
                struct_total["tables"] += s["tables"]
                struct_total["images"] += s["images"]

            st.dataframe(stats_data, use_container_width=True)

            st.markdown("### Document Health")
            dh1, dh2 = st.columns([1,1])
            with dh1:
                ocr_score = _compute_ocr_quality("\n\n".join(combined_texts))
                fig_g = go.Figure(go.Indicator(
                    mode="gauge+number",
                    value=ocr_score,
                    number={'suffix': " /100"},
                    title={'text': "OCR Quality Score"},
                    gauge={'axis': {'range': [0,100]},
                           'bar': {'thickness': 0.3},
                           'steps': [
                               {'range': [0,50], 'color': "#fce4ec"},
                               {'range': [50,75], 'color': "#fff3e0"},
                               {'range': [75,100], 'color': "#e8f5e9"},
                           ]}
                ))
                st.plotly_chart(fig_g, use_container_width=True, key="ocr_gauge_chart")

            with dh2:
                labels = ["Text", "Tables", "Images"]
                values = [max(1, struct_total["text"]), max(1, struct_total["tables"]), max(1, struct_total["images"])]
                fig_donut = px.pie(values=values, names=labels, hole=0.6, title="Document Structure Overview")
                st.plotly_chart(fig_donut, use_container_width=True, key="structure_donut_chart")

            secs = _extract_sections("\n\n".join(combined_texts))
            sec_labels, cat_labels, mat = _missing_matrix(secs)
            fig_heat = px.imshow(mat, aspect="auto", color_continuous_scale="Viridis",
                                 labels=dict(x="Section", y="Missing Type", color="Count"),
                                 x=sec_labels, y=cat_labels)
            fig_heat.update_layout(title="Missing Data Heatmap")
            st.plotly_chart(fig_heat, use_container_width=True, key="missing_data_heatmap")

        # ---------- Chat history ----------
        for i, m in enumerate(st.session_state.chat_history):
            role = "You" if m["role"] == "user" else "Assistant"
            st.markdown(f"**{role}:** {m['content']}")
            
            # Display chart if available
            if m.get("chart") and m["role"] == "assistant":
                st.markdown("📊 **Visualization:**")
                st.plotly_chart(m["chart"], use_container_width=True, key=f"chart_{i}_{hash(str(m.get('chart_message', '')))}")
                
                # Show chart info
                if m.get("chart_message"):
                    st.info(m["chart_message"])

        # ---------- Q&A input ----------
        with st.form("qa_form_docs", clear_on_submit=True):
            user_q = st.text_input(
                "Your question (can ask about specific documents or compare them):",
                key="qa_input"
            )
            submitted = st.form_submit_button("Ask")

        if submitted and user_q:
            st.session_state.chat_history.append({"role": "user", "content": user_q})
            with st.spinner("Generating response..."):
                ans = None
                chart = None
                chart_message = ""
                
                try:
                    docs_to_use = select_docs_for_query(user_q, st.session_state.documents)
                    
                    # Check for visualization intent first
                    viz_success, viz_msg, viz_chart = generate_visualization(user_q, docs_to_use)
                    if viz_success:
                        chart = viz_chart
                        chart_message = viz_msg
                        if ANSWER_LANGUAGE == "Bahasa Indonesia":
                            ans = f"✅ {chart_message}\n\nSaya telah membuat visualisasi berdasarkan data dari dokumen Anda. Grafik ditampilkan di bawah ini."
                        else:
                            ans = f"✅ {chart_message}\n\nI have created a visualization based on your document data. The chart is displayed below."
                    else:
                        # Try average calculation
                        avg_results = compute_average_per_doc(user_q, docs_to_use)
                        if avg_results:
                            if len(avg_results) == 1:
                                r = avg_results[0]
                                val_str = f"{r['value']:,.4f}".replace(",", "X").replace(".", ",").replace("X", ".")
                                if ANSWER_LANGUAGE == "Bahasa Indonesia":
                                    ans = f"Rata-rata di dokumen **{r['doc']}** untuk kolom **{r['column']}** adalah **{val_str}** (n={r['n']})."
                                else:
                                    ans = f"The average in document **{r['doc']}** for column **{r['column']}** is **{val_str}** (n={r['n']})."
                            else:
                                ans_lines = []
                                for r in avg_results:
                                    val_str = f"{r['value']:,.4f}".replace(",", "X").replace(".", ",").replace("X", ".")
                                    ans_lines.append(f"- **{r['doc']}** → {val_str} (n={r['n']})")
                                best = max(avg_results, key=lambda x: x["value"])
                                if ANSWER_LANGUAGE == "Bahasa Indonesia":
                                    ans = "Hasil rata-rata per dokumen:\n" + "\n".join(ans_lines)
                                    ans += f"\n\n📊 Nilai tertinggi ada di **{best['doc']}** dengan rata-rata {best['value']:.2f}."
                                else:
                                    ans = "Average per document:\n" + "\n".join(ans_lines)
                                    ans += f"\n\n📊 Highest is in **{best['doc']}** with average {best['value']:.2f}."
                
                except Exception as e:
                    st.error(f"Error processing query: {str(e)}")
                    avg_results = []
                
                if not ans:
                    if not google_api_key:
                        ans = "Please provide a valid Google API Key."
                    else:
                        ans = generate_response_with_fallback(st.session_state.ocr_content, user_q)
                
                # Store response with chart info
                response_data = {"role": "assistant", "content": ans}
                if chart:
                    response_data["chart"] = chart
                    response_data["chart_message"] = chart_message
                
                st.session_state.chat_history.append(response_data)
                st.rerun()

    else:
        st.info("No documents processed yet. You can either upload files or just type a URL below and press Ask.")
        with st.form("qa_form_nodocs", clear_on_submit=True):
            user_q = st.text_input(
                "Your question (can ask about specific documents or compare them):",
                key="qa_input"
            )
            submitted = st.form_submit_button("Ask")

        if submitted and user_q:
            url_candidate = st.session_state.get("url_input")
            if url_candidate and not st.session_state.get("ocr_content"):
                with st.spinner("Processing URL before answering..."):
                    success, msg = process_url_to_content(url_candidate)
                    if not success:
                        st.error(msg)
                        st.stop()
            if st.session_state.get("ocr_content"):
                st.session_state.chat_history.append({"role": "user", "content": user_q})
                with st.spinner("Generating response..."):
                    ans = None
                    chart = None
                    chart_message = ""
                    
                    try:
                        docs_to_use = select_docs_for_query(user_q, st.session_state.documents)
                        
                        # Check for visualization intent first
                        viz_success, viz_msg, viz_chart = generate_visualization(user_q, docs_to_use)
                        if viz_success:
                            chart = viz_chart
                            chart_message = viz_msg
                            if ANSWER_LANGUAGE == "Bahasa Indonesia":
                                ans = f"✅ {chart_message}\n\nSaya telah membuat visualisasi berdasarkan data dari dokumen Anda. Grafik ditampilkan di bawah ini."
                            else:
                                ans = f"✅ {chart_message}\n\nI have created a visualization based on your document data. The chart is displayed below."
                        else:
                            # Try average calculation
                            avg_results = compute_average_per_doc(user_q, docs_to_use)
                            if avg_results:
                                if len(avg_results) == 1:
                                    r = avg_results[0]
                                    val_str = f"{r['value']:,.4f}".replace(",", "X").replace(".", ",").replace("X", ".")
                                    if ANSWER_LANGUAGE == "Bahasa Indonesia":
                                        ans = f"Rata-rata di dokumen **{r['doc']}** untuk kolom **{r['column']}** adalah **{val_str}** (n={r['n']})."
                                    else:
                                        ans = f"The average in document **{r['doc']}** for column **{r['column']}** is **{val_str}** (n={r['n']})."
                                else:
                                    ans_lines = []
                                    for r in avg_results:
                                        val_str = f"{r['value']:,.4f}".replace(",", "X").replace(".", ",").replace("X", ".")
                                        ans_lines.append(f"- **{r['doc']}** → {val_str} (n={r['n']})")
                                    best = max(avg_results, key=lambda x: x["value"])
                                    if ANSWER_LANGUAGE == "Bahasa Indonesia":
                                        ans = "Hasil rata-rata per dokumen:\n" + "\n".join(ans_lines)
                                        ans += f"\n\n📊 Nilai tertinggi ada di **{best['doc']}** dengan rata-rata {best['value']:.2f}."
                                    else:
                                        ans = "Average per document:\n" + "\n".join(ans_lines)
                                        ans += f"\n\n📊 Highest is in **{best['doc']}** with average {best['value']:.2f}."
                    
                    except Exception as e:
                        st.error(f"Error processing query: {str(e)}")
                        avg_results = []
                    
                    if not ans:
                        if not google_api_key:
                            ans = "Please provide a valid Google API Key."
                        else:
                            ans = generate_response_with_fallback(st.session_state.ocr_content, user_q)
                    
                    # Store response with chart info
                    response_data = {"role": "assistant", "content": ans}
                    if chart:
                        response_data["chart"] = chart
                        response_data["chart_message"] = chart_message
                    
                    st.session_state.chat_history.append(response_data)
                    st.rerun()
            else:
                st.warning("Please provide a URL or upload a document first.")

# Tampilkan konten hasil OCR/ekstraksi
if st.session_state.get("documents"):
    with st.expander("📄 View All Document Contents"):
        for i, doc in enumerate(st.session_state.documents):
            st.markdown(f"### {doc['name']} ({doc['type']})")
            if isinstance(doc['content'], str) and ("<table" in doc['content'] or "<p" in doc['content'] or "</table>" in doc['content']):
                st.markdown(doc['content'], unsafe_allow_html=True)
            else:
                st.markdown(doc['content'])
            st.markdown("---")
